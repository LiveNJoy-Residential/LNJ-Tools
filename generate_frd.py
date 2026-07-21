#!/usr/bin/env python3
"""
Convert FRD.md to FRD.docx using python-docx.
Produces a professionally formatted Word / Google Docs-compatible document.
"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE     = '1F4E79'
MID_BLUE      = '2E74B5'
ALT_ROW       = 'D6E4F0'
NOTE_BG       = 'FFF2CC'
DARK_BLUE_RGB = RGBColor(0x1F, 0x4E, 0x79)
MID_BLUE_RGB  = RGBColor(0x2E, 0x74, 0xB5)
WHITE_RGB     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_RGB      = RGBColor(0x44, 0x44, 0x44)

# ── Low-level XML helpers ──────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_bottom_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom    = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),  'single')
    bottom.set(qn('w:sz'),   '4')
    bottom.set(qn('w:color'),'CCCCCC')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


# ── Markdown inline-formatting parser ─────────────────────────────────────────

def add_inline(para, text: str, size: Pt = Pt(10)):
    """Split text on **bold** and `code` markers and add correctly styled runs."""
    token_re = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.size = size
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(max(size.pt - 1, 8))
        else:
            r = para.add_run(part)
            r.font.size = size
    return para


def strip_md(text: str) -> str:
    """Remove all markdown inline markers, returning plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`',        r'\1', text)
    text = re.sub(r'\*(.+?)\*',      r'\1', text)
    return text.strip()


# ── Table builder ──────────────────────────────────────────────────────────────

def _col_widths(headers):
    n   = len(headers)
    h0  = (headers[0] if headers else '').strip().lower()
    if n == 2:
        if h0 in ('id',):
            return [Inches(0.75), Inches(5.5)]
        if 'term' in h0:
            return [Inches(1.6), Inches(4.65)]
        if 'data source' in h0 or 'source' in h0:
            return [Inches(1.8), Inches(4.45)]
        return [Inches(2.0), Inches(4.25)]
    if n == 3:
        if 'data source' in h0 or 'source' in h0:
            return [Inches(1.9), Inches(2.3), Inches(2.05)]
        if 'tab' in h0:
            return [Inches(1.8), Inches(4.45)]
        if 'actor' in h0:
            return [Inches(1.5), Inches(2.1), Inches(2.65)]
        return [Inches(1.4), Inches(2.2), Inches(2.65)]
    if n == 4:
        return [Inches(0.6), Inches(2.85), Inches(0.8), Inches(1.95)]
    w = 6.2 / n
    return [Inches(w)] * n


def build_table(doc, table_lines):
    """Parse markdown table lines and insert a formatted Word table."""
    parsed = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        parsed.append(cells)

    # Drop separator rows
    def is_sep(row):
        return all(re.fullmatch(r'[-:\s]+', c) for c in row if c.strip())
    data = [r for r in parsed if not is_sep(r)]
    if len(data) < 2:
        return

    headers = data[0]
    rows    = data[1:]
    ncols   = len(headers)

    table = doc.add_table(rows=1, cols=ncols)
    table.style = 'Table Grid'

    # ── Header row ──
    hdr = table.rows[0]
    for ci, h in enumerate(headers):
        if ci >= ncols:
            break
        cell = hdr.cells[ci]
        cell.text = ''
        set_cell_bg(cell, DARK_BLUE)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r   = p.add_run(strip_md(h))
        r.bold = True
        r.font.color.rgb = WHITE_RGB
        r.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # ── Data rows ──
    for ri, row in enumerate(rows):
        tr = table.add_row()
        bg = ALT_ROW if ri % 2 == 1 else None
        for ci, txt in enumerate(row):
            if ci >= ncols:
                break
            cell = tr.cells[ci]
            cell.text = ''
            if bg:
                set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_inline(p, txt, size=Pt(8.5))

    # ── Column widths ──
    widths = _col_widths(headers)
    for row in table.rows:
        for ci, w in enumerate(widths):
            if ci < len(row.cells):
                try:
                    row.cells[ci].width = w
                except Exception:
                    pass

    doc.add_paragraph()   # breathing room after every table


# ── Cover page ────────────────────────────────────────────────────────────────

def build_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Functional Requirements Document')
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE_RGB

    # Product name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LNJ Audit Bot')
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE_RGB

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LiveNjoy Residential  ·  ResMan Audit Automation')
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY_RGB

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Metadata table
    mt = doc.add_table(rows=5, cols=2)
    mt.style = 'Table Grid'
    meta = [
        ('Version',  '2.0'),
        ('Status',   'Final'),
        ('Authors',  'Concession Audit Engine   |   Revenue Integrity Engine'),
        ('Company',  'LiveNjoy Residential'),
        ('Platform', 'ResMan Property Management'),
    ]
    for i, (k, v) in enumerate(meta):
        lc = mt.rows[i].cells[0]
        vc = mt.rows[i].cells[1]
        set_cell_bg(lc, MID_BLUE)
        lc.text = ''
        p = lc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(k)
        r.bold = True
        r.font.color.rgb = WHITE_RGB
        r.font.size = Pt(10)
        vc.text = ''
        p2 = vc.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.size = Pt(10)

    for row in mt.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(4.7)

    doc.add_page_break()


# ── Section-level page-break helper ───────────────────────────────────────────

def _page_break_before(paragraph):
    """Insert a page break before the given paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()
    pb  = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)


# ── Main converter ────────────────────────────────────────────────────────────

def convert(md_path: str, docx_path: str):
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.0)

    build_cover(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Skip the cover-page header block (everything up to and including the first ---)
    start = 0
    for idx, line in enumerate(lines):
        if line.strip() == '---':
            start = idx + 1
            break

    i          = start
    table_buf  = []
    in_table   = False

    # Track whether we're inside the top-level ## heading
    # so we can inject page-breaks before each new ## section
    prev_h1_para = None

    while i < len(lines):
        line    = lines[i].rstrip('\n')
        stripped = line.strip()

        # ── Table accumulation ──────────────────────────────────────
        if stripped.startswith('|'):
            in_table = True
            table_buf.append(stripped)
            i += 1
            continue

        # ── Flush pending table ─────────────────────────────────────
        if in_table:
            in_table = False
            build_table(doc, table_buf)
            table_buf = []
            # fall through to process current (non-table) line

        # ── Headings ────────────────────────────────────────────────
        if stripped.startswith('#### '):
            h = doc.add_heading(stripped[5:], level=3)
            for run in h.runs:
                run.font.color.rgb = MID_BLUE_RGB

        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=2)
            for run in h.runs:
                run.font.color.rgb = MID_BLUE_RGB

        elif stripped.startswith('## '):
            title = stripped[3:]
            h = doc.add_heading(title, level=1)
            for run in h.runs:
                run.font.color.rgb = DARK_BLUE_RGB
            # Page-break before every top-level section (except the very first)
            if prev_h1_para is not None:
                _page_break_before(h)
            prev_h1_para = h

        elif stripped.startswith('# '):
            # Document title — already on cover page, skip
            pass

        # ── Bullet lists ─────────────────────────────────────────────
        elif re.match(r'^  [-*] ', stripped) or re.match(r'^    [-*] ', stripped):
            # indented sub-bullet
            p = doc.add_paragraph(style='List Bullet 2')
            add_inline(p, re.sub(r'^[\s\-\*]+', '', stripped), size=Pt(10))

        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, stripped[2:], size=Pt(10))

        # ── Block-quote (note / severity key) ───────────────────────
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            # light yellow shading via XML — attach to the run's paragraph
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  NOTE_BG)
            pPr.append(shd)
            add_inline(p, stripped[2:], size=Pt(9.5))
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # ── Horizontal rule ──────────────────────────────────────────
        elif stripped == '---':
            pass  # Skip — visual separation is handled by page breaks & spacing

        # ── Empty line ───────────────────────────────────────────────
        elif stripped == '':
            pass  # suppress excessive whitespace

        # ── Regular paragraph ────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.space_before = Pt(2)
            add_inline(p, stripped, size=Pt(10))

        i += 1

    # Flush any table still open at EOF
    if in_table and table_buf:
        build_table(doc, table_buf)

    doc.save(docx_path)
    print(f'[OK] Saved → {docx_path}')


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    base     = os.path.dirname(os.path.abspath(__file__))
    md_path  = os.path.join(base, 'FRD.md')
    out_path = os.path.join(base, 'FRD.docx')
    convert(md_path, out_path)
